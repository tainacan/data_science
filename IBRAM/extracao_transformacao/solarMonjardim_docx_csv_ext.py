from docx.api import Document
from collections import defaultdict
import win32com.client as win32
from time import sleep #define um periodo de tempo entre uma ação e outra
from tqdm import tqdm #adiciona uma barra de "loading"
import glob
import csv
import os
import pandas as pd
import re

resultado = pd.DataFrame()
path = ''

def re_match(match, texto):
    
    try :
        return re.search(match, texto)[0].strip()
    except:
        return ""

def splitTxt(txt):
    if ":" in txt:
        return txt.split(":")[1].strip()
    else:
        return txt.split("\n")[1].strip()
#%%
for infile in tqdm(glob.glob(os.path.join(path, '*.docx') ), desc="processando dados", unit="files"):
    print("Verificando o arquivo {}".format(infile.split("\\")[-1]))
    
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = 0
    
    doc = Document(infile)
        
    resultDict = {}

    doc_fot = "(?<=Documentação fotográfica: )((.|\n)*)(?=Foto n°:)"
    foto = "(?<=Foto n°: \n)((.|\n)*)(?=Filme:)"
    filme = "(?<=Filme:)((.|\n)*)(?=Negativo:)"
    negativo = "(?<=Negativo: )((.|\n)*)(?=Fotográfico:)"
    fotografico = "(?<=Fotográfico: )((.|\n)*)(?=Data:)"
    foto_data = "(?<=Data: )((.|\n)*)(?=)"
    pT_resp = "(?<=Resp.: )((.|\n)*)(?=Data:)"
    pT_data = "(?<=Data: )((.|\n)*)(?=)"
    res_restauradores = "(?<=Restauradores: )((.|\n)*)(?=Data:)"
    res_data = "(?<=Data: )((.|\n)*)(?=)"
    revT_resp = "(?<=Resp.: )((.|\n)*)(?=Data:)"
    revT_data = "(?<=Data: )((.|\n)*)(?=)"
    dim_altura = "(?<=altura: )((.|\n)*)(?=largura:)"
    dim_largura = "(?<=largura: )((.|\n)*)(?=profundidade:)"
    dim_profundidade = "(?<=profundidade: )((.|\n)*)(?=)"
    
    for pL_tipo in doc.tables[0].cell(13,0).text.strip().split("\n")[1:]:
        if "X" in pL_tipo.strip():
            protLegalTipo = pL_tipo.strip().split(") ")[1]
            
    for condSeg in doc.tables[0].cell(14,0).text.strip().split("\n")[2].split("( "):
        if "X" in condSeg.strip():
            condSegValor = condSeg.strip().split(") ")[1].strip()
            
    for estCons in " ".join(doc.tables[0].cell(15,0).text.strip().split("\n")[2:]).split("( "):
        if "X" in estCons.strip():
            estConsValor = estCons.strip().split(") ")[1].strip()
    
    refBib = []
    for bibliografia in doc.tables[0].cell(22,0).text.strip().split("\n")[1: ]:
        refBib.append(bibliografia.strip())
    refBib = list(filter(None, refBib))
    
    
    resultDict['Arquivo'] = infile.split("\\")[-1]
    resultDict['UF — Município'] = splitTxt(doc.tables[0].cell(1,0).text.strip())
    resultDict['Distrito'] = splitTxt(doc.tables[0].cell(2,0).text.strip())
    resultDict['Endereço'] = splitTxt(doc.tables[0].cell(3,0).text.strip())
    resultDict['Acervo'] = splitTxt(doc.tables[0].cell(4,0).text.strip())
    resultDict['Local no prédio'] = splitTxt(doc.tables[0].cell(5,0).text.strip())
    resultDict['Proprietário'] = splitTxt(doc.tables[0].cell(6,0).text.strip())
    resultDict['Responsável imediato'] = splitTxt(doc.tables[0].cell(7,0).text.strip())
    resultDict['FOTO_Desc'] = doc.tables[0].cell(8,0).text.strip()
    resultDict['Marcas/Inscrições/Legendas'] = splitTxt(doc.tables[0].cell(9,0).text.strip())
    resultDict['Documentação fotográfica'] = re_match(doc_fot, doc.tables[0].cell(10,0).text.strip())
    resultDict['Foto n°'] = "||".join(re_match(foto, doc.tables[0].cell(10,0).text.strip()).split("\n"))
    resultDict['Filme'] = re_match(filme, doc.tables[0].cell(10,0).text.strip())
    resultDict['Negativo'] = re_match(negativo, doc.tables[0].cell(10,0).text.strip())
    resultDict['Fotográfico'] = re_match(fotografico, doc.tables[0].cell(10,0).text.strip())
    resultDict['Foto_Data'] = re_match(foto_data, doc.tables[0].cell(10,0).text.strip())
    resultDict['Proteção'] = splitTxt(doc.tables[0].cell(11,0).text.strip())
    resultDict['Descrição'] = splitTxt(doc.tables[0].cell(12,0).text.strip())
    resultDict['Proteção legal_Documento'] = splitTxt(doc.tables[0].cell(13,0).text.strip().split("\n")[0])
    resultDict['Proteção legal_Tipo'] = protLegalTipo
    resultDict['Condições de segurança'] = condSegValor
    resultDict['Estado de conservação'] = estConsValor
    resultDict['Especificação do estado de conservação'] = splitTxt(doc.tables[0].cell(16,0).text.strip())
    resultDict['Restaurações'] = splitTxt(doc.tables[0].cell(17,0).text.strip())
    resultDict['Características técnicas'] = splitTxt(doc.tables[0].cell(18,0).text.strip())
    resultDict['Características estilísticas'] = splitTxt(doc.tables[0].cell(19,0).text.strip())
    resultDict['Características iconográficas/ornamentais'] = splitTxt(doc.tables[0].cell(20,0).text.strip())
    resultDict['Dados históricos'] = splitTxt(doc.tables[0].cell(21,0).text.strip())
    resultDict['Referências bibliográficas/arquivísticas'] = "||".join(refBib)
    resultDict['Observações'] = splitTxt(doc.tables[0].cell(23,0).text.strip())
    resultDict['Preenchimento técnico_Responsável'] = re_match(pT_resp, doc.tables[0].cell(24,0).text.strip())
    resultDict['Preenchimento técnico_Data'] = re_match(pT_data, doc.tables[0].cell(24,0).text.strip())
    resultDict['Dados complementares'] = splitTxt(doc.tables[0].cell(25,0).text.strip())
    resultDict['Título'] = splitTxt(doc.tables[0].cell(1,1).text.strip())
    resultDict['Número'] = splitTxt(doc.tables[0].cell(1,4).text.strip())
    resultDict['Objeto'] = splitTxt(doc.tables[0].cell(2,1).text.strip())
    resultDict['Classe'] = splitTxt(doc.tables[0].cell(3,1).text.strip())
    resultDict['Subclasse'] = splitTxt(doc.tables[0].cell(4,1).text.strip())
    resultDict['Datação'] = splitTxt(doc.tables[0].cell(5,1).text.strip())
    resultDict['Autoria'] = splitTxt(doc.tables[0].cell(6,1).text.strip())
    resultDict['Material/Técnica'] = splitTxt(doc.tables[0].cell(7,1).text.strip())
    resultDict['Restauradores'] = re_match(res_restauradores, doc.tables[0].cell(17,2).text.strip())
    resultDict['Restaurações_Data'] = re_match(res_data, doc.tables[0].cell(17,2).text.strip())
    resultDict['Revisão Técnica_Responsável'] = re_match(revT_resp, doc.tables[0].cell(24,2).text.strip())
    resultDict['Revisão Técnica_Data'] = re_match(revT_data, doc.tables[0].cell(24,2).text.strip())
    resultDict['N° anterior — tombo n°'] = splitTxt(doc.tables[0].cell(2,4).text.strip())
    resultDict['Origem'] = splitTxt(doc.tables[0].cell(3,4).text.strip())
    resultDict['Procedência'] =splitTxt(doc.tables[0].cell(4,4).text.strip())
    resultDict['Modo de aquisição/Data'] = splitTxt(doc.tables[0].cell(5,4).text.strip())
    resultDict['Conjunto com n° — peças relacionadas'] = splitTxt(doc.tables[0].cell(6,4).text.strip())
    resultDict['Termos de indexação'] = splitTxt(doc.tables[0].cell(7,4).text.strip())
    resultDict['Dimensões (cm)_altura'] = re_match(dim_altura, doc.tables[0].cell(10,4).text.strip().lower())
    resultDict['Dimensões (cm)_largura'] = re_match(dim_largura, doc.tables[0].cell(10,4).text.strip().lower())
    resultDict['Dimensões (cm)_profundidade'] = re_match(dim_profundidade, doc.tables[0].cell(10,4).text.strip().lower())

    sleep(5)
    resultado = resultado.append(resultDict, ignore_index=True)
word.Quit()

#%%
for key, value in resultDict.items():
    print(key, " - ", value)
#%%
resultado.to_csv("", index = False)